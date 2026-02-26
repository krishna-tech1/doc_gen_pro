"""
services/document_service.py — python-docx document generation.

Builds three document types:
  - Circular
  - Proposal
  - Event Report

All documents share:
  - Institutional logo header
  - Consistent styling (fonts, spacing, borders)
  - AI-generated content sections
Returns the absolute path to the generated .docx file.
"""

from __future__ import annotations
import os
import uuid
from typing import Optional, List
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ─────────────────────────────────────────────────────────────────
LOGO_PATH      = os.path.join(os.path.dirname(__file__), "..", "logo.png")
BANNER_PNG     = os.path.join(os.path.dirname(__file__), "..", "banner.png")
BANNER_JPG     = os.path.join(os.path.dirname(__file__), "..", "banner.jpg")
OUTPUT_DIR     = os.path.join(os.path.dirname(__file__), "..", "generated")
INSTITUTION    = "AVICHI COLLEGE OF ARTS AND SCIENCE"
HEADING_COLOR  = RGBColor(0x6A, 0x15, 0x18)   # institutional maroon
ACCENT_COLOR   = RGBColor(0x00, 0x00, 0x00)   # black


def _ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _unique_filename(prefix: str) -> str:
    _ensure_output_dir()
    filename = f"{prefix}_{uuid.uuid4().hex[:8]}.docx"
    return os.path.join(OUTPUT_DIR, filename)


# ── Styling helpers ────────────────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Helper to add borders around a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag  = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),  kwargs.get(edge, {}).get("val",  "single"))
        tag.set(qn("w:sz"),   kwargs.get(edge, {}).get("sz",   "4"))
        tag.set(qn("w:color"),kwargs.get(edge, {}).get("color","1A5376"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _add_logo_header(doc: Document, subtitle: str = ""):
    """Insert banner or logo + institution name as the document header."""
    # Check if a banner image exists (PNG or JPG)
    banner_path = None
    
    # Try all possible names the user might have saved the header as
    possible_banners = [
        os.path.join(os.path.dirname(__file__), "..", "banner.png"),
        os.path.join(os.path.dirname(__file__), "..", "banner.jpg"),
        os.path.join(os.path.dirname(__file__), "..", "logo.png"),
        os.path.join(os.path.dirname(__file__), "..", "logo.jpeg"),
        os.path.join(os.path.dirname(__file__), "..", "logo.jpg")
    ]
    
    for path in possible_banners:
        if os.path.exists(path):
            banner_path = path
            break

    # If banner image exists, use it spanning the full width
    if banner_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(banner_path, width=Inches(6.5))
        return


    # Fallback to text header if banner isn't found
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(4.3)
    table.columns[2].width = Inches(1.2)

    # Logo cell
    logo_cell = table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Inches(1.0))
    else:
        run.text = "🏛"
        run.font.size = Pt(36)

    # Text cell
    text_cell = table.cell(0, 1)
    p1 = text_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(INSTITUTION)
    r1.bold       = True
    r1.font.size  = Pt(14)
    r1.font.color.rgb = HEADING_COLOR

    p2 = text_cell.add_paragraph("(A Co-Educational Institution Affiliated to the University of Madras)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(9)

    p3 = text_cell.add_paragraph("130A, Arcot Road, Virugambakkam, Chennai - 600 092.\nPh : +91 44 2376 4227 | E.Mail : avichicollege2017@gmail.com")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].font.size = Pt(8)

    # Blank right cell to balance the header or add a second logo if available
    right_cell = table.cell(0, 2)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # spacer


def _add_section(doc: Document, heading: str, body: str):
    """Add a titled section with consistent heading + body paragraph."""
    h = doc.add_heading(heading, level=2)
    h.runs[0].font.color.rgb = HEADING_COLOR
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph(body)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(11)


def _add_divider(doc: Document):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), "1A5376")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Circular ──────────────────────────────────────────────────────────────────

def generate_circular(data: dict, ai_content: dict = None) -> str:
    """
    Build a formatted Circular .docx matching the user's specific template.
    data keys: event_name, date, end_date, time, department, chief_guest, description
    """
    doc  = Document()
    _set_page_margins(doc)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    _add_logo_header(doc, "")

    # Title
    t_para = doc.add_paragraph()
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = data.get('document_title', 'CIRCULAR')
    run_t = t_para.add_run(title_text)
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(14)
    run_t.bold = True
    run_t.underline = True

    doc.add_paragraph()

    # Date Align Right
    d_para = doc.add_paragraph()
    d_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_line = data.get('date_line') or f"Date: {data.get('date', '')}"
    d_para.add_run(date_line).font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Main Body Text
    description = data.get("description")
    if description and description.strip():
        # User provided custom body - use ONLY this
        p_body = doc.add_paragraph(description)
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        # Paragraph 1
        p_body = doc.add_paragraph()
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if data.get('circular_body_1'):
             p_body.add_run(data['circular_body_1'])
        else:
            dept = data.get("department", "...")
            event = data.get("event_name", "...")
            time_val = data.get("time") or "___ AM/PM"
            date_val = data.get("date", "")
            p_body.add_run("This is to inform that the ")
            p_body.add_run(f"Department of {dept}").bold = True
            p_body.add_run(" will be organizing ")
            p_body.add_run(f"{event}").bold = True
            p_body.add_run(f" on {date_val} at {time_val} in the college premises.")
        
        doc.add_paragraph()
        
        # Paragraph 2
        p_aim = doc.add_paragraph()
        p_aim.paragraph_format.line_spacing = 1.15
        p_aim.add_run(data.get('circular_body_2', "The program aims to enrich students with knowledge and practical exposure in the relevant field."))
        
        doc.add_paragraph()
        
        # Paragraph 3
        p_guest = doc.add_paragraph()
        p_guest.paragraph_format.line_spacing = 1.15
        if data.get('circular_body_3'):
            p_guest.add_run(data['circular_body_3'])
        else:
            p_guest.add_run("We are honored to have ")
            p_guest.add_run(data.get("chief_guest", "...")).bold = True
            p_guest.add_run(" as the Chief Guest for this event.")
        
        doc.add_paragraph()
        
        # Paragraph 4
        p_note = doc.add_paragraph()
        p_note.paragraph_format.line_spacing = 1.15
        p_note.add_run(data.get('circular_body_4', "All concerned are requested to take note of the same and participate actively."))

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    p_sig = doc.add_paragraph()
    p_sig.add_run(data.get('footer_left', "Event Coordinator") + "\t\t\t\t\t" + data.get('footer_right', "Principal")).bold = True

    doc.add_paragraph()
    
    # Copy To
    p_copy = doc.add_paragraph()
    p_copy.add_run(data.get('copy_to_label', "Copy To:")).bold = True
    
    if data.get('copy_to_list'):
        doc.add_paragraph(data['copy_to_list']).paragraph_format.left_indent = Inches(0.5)
    else:
        p_h = doc.add_paragraph("• All HODs")
        p_h.paragraph_format.left_indent = Inches(0.5)
        p_iq = doc.add_paragraph("• IQAC")
        p_iq.paragraph_format.left_indent = Inches(0.5)

    path = _unique_filename("circular")
    doc.save(path)
    return path

    path = _unique_filename("circular")
    doc.save(path)
    return path


# ── Proposal ──────────────────────────────────────────────────────────────────

def generate_proposal(data: dict, ai_content: dict = None) -> str:
    """
    Build a structured Permission Letter Proposal .docx based on user template.
    """
    doc = Document()
    _set_page_margins(doc)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Date
    p_date = doc.add_paragraph()
    date_line = data.get('proposal_date_line') or f"Date: {data.get('proposal_date', '')}"
    p_date.add_run(date_line).font.size = Pt(12)
    doc.add_paragraph()

    # From Block
    p_from_label = doc.add_paragraph()
    p_from_label.add_run(data.get('from_label', "From")).bold = True
    
    p_from = doc.add_paragraph()
    p_from.add_run(data.get('from_name_line', data.get('from_name', '')) + "\n").font.size = Pt(12)
    p_from.add_run(data.get('from_designation_line', data.get('from_designation', '')) + "\n").font.size = Pt(12)
    p_from.add_run(data.get('department_line', data.get('department', '')) + "\n").font.size = Pt(12)
    p_from.add_run(data.get('college_name_from', "Avichi College of Arts and Science") + "\n").font.size = Pt(12)
    p_from.add_run(data.get('college_address_from', "Virugambakkam-92")).font.size = Pt(12)
    
    doc.add_paragraph()

    # To Block
    p_to_label = doc.add_paragraph()
    p_to_label.add_run(data.get('to_label', "To")).bold = True
    
    p_to = doc.add_paragraph()
    p_to.add_run(data.get('to_recipient', "The Principal") + "\n").font.size = Pt(12)
    p_to.add_run(data.get('college_name_to', "Avichi College of Arts and Science") + "\n").font.size = Pt(12)
    p_to.add_run(data.get('college_address_to', "Virugambakkam-92")).font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(data.get('salutation', "Respected Madam,"))
    doc.add_paragraph()

    # Subject
    p_sub = doc.add_paragraph()
    sub_line = data.get('subject_line') or f"Subject: Permission Letter for {data.get('event_name', '')}"
    p_sub.add_run(sub_line).bold = True
    
    doc.add_paragraph()

    # Body Paragraph 1
    p_body1 = doc.add_paragraph()
    p_body1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if data.get('proposal_body_1'):
        p_body1.add_run(data['proposal_body_1'])
    else:
        p_body1.add_run(f"I propose to conduct a session on ")
        p_body1.add_run(f"{data.get('event_topic', '')}").bold = True
        p_body1.add_run(f" for ")
        p_body1.add_run(f"{data.get('target_audience', '')}").bold = True
        p_body1.add_run(f" on {data.get('event_date', '')} at {data.get('event_time', '')} in ")
        p_body1.add_run(f"{data.get('venue', '')}").bold = True
        p_body1.add_run(". ")
        p_body1.add_run(f"We have invited ")
        p_body1.add_run(f"{data.get('resource_person', '')}").bold = True
        p_body1.add_run(" as the resource person.")

    doc.add_paragraph()

    # Body Paragraph 2
    p_body2 = doc.add_paragraph()
    p_body2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if data.get('proposal_body_2'):
        p_body2.add_run(data['proposal_body_2'])
    else:
        desc = data.get('short_description', '...')
        p_body2.add_run(f"The session will focus on ")
        p_body2.add_run(f"{desc}").bold = True
        p_body2.add_run(". It aims to provide practical exposure and enhance knowledge among students.")

    doc.add_paragraph()

    doc.add_paragraph(data.get('request_sentence', "I kindly request your permission to proceed with the arrangements for this event."))
    doc.add_paragraph(data.get('thanks_sentence', "Thank you for your support and cooperation."))

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing
    p_closing = doc.add_paragraph()
    p_closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing_label = data.get('closing_label', "Yours sincerely,")
    p_closing.add_run(f"{closing_label}\n\n").font.size = Pt(12)
    p_closing.add_run(f"{data.get('from_name', '')}").bold = True

    path = _unique_filename("proposal")
    doc.save(path)
    return path


# ── Event Report ──────────────────────────────────────────────────────────────

def generate_report(data: dict, ai_content: dict, images: List[str],
                    sdg_description: Optional[str] = None) -> str:
    """
    Build a formal Event Report .docx.
    data keys: event_title, summary, num_participants, sdg_number, location_name, report_header, sdg_goal
    """
    doc = Document()
    _set_page_margins(doc)
    
    # Header
    header_text = data.get('report_header', "EVENT REPORT")
    _add_logo_header(doc, header_text)
    _add_divider(doc)

    # Title
    title_text = (data.get("event_title") or "Event Title").upper()
    t = doc.add_heading(title_text, level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = HEADING_COLOR

    doc.add_paragraph()

    # Summary table
    sdg_display = data.get('sdg_goal') or f"SDG {data.get('sdg_number', '')} – {data.get('sdg_name', '')}"
    meta = [
        ("Number of Participants",  str(data.get("num_participants", ""))),
        ("Event Date",              data.get("date") or datetime.now().strftime("%d %B %Y")),
        ("Location",                data.get("location_name", "On Campus")),
        ("SDG Goal",                sdg_display),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = value
        tbl.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    _add_divider(doc)
    doc.add_paragraph()

    _add_section(doc, "Introduction",         ai_content.get("introduction", ""))
    _add_section(doc, "Objectives",           ai_content.get("objectives", ""))
    _add_section(doc, "Proceedings Summary",  ai_content.get("outcome", ""))
    _add_section(doc, "SDG Alignment",
                 data.get('sdg_alignment') or sdg_description or ai_content.get("sdg_alignment") or ai_content.get("conclusion", ""))
    _add_section(doc, "Conclusion",           ai_content.get("conclusion", ""))

    # Images
    if images:
        doc.add_heading("Event Gallery", level=2).runs[0].font.color.rgb = HEADING_COLOR
        for img_path in images:
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(img_path, width=Inches(4.5))
                except Exception:
                    pass

    path = _unique_filename("report")
    doc.save(path)
    return path


# ── Page margin helper ────────────────────────────────────────────────────────

def _set_page_margins(doc: Document, margin: float = 1.0):
    """Set uniform 1-inch page margins."""
    for section in doc.sections:
        section.top_margin    = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin   = Inches(margin)
        section.right_margin  = Inches(margin)
